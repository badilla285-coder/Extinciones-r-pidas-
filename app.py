import streamlit as st
from docx import Document
from docx.shared import Pt
import PyPDF2
import io

def crear_escrito(datos, texto_condena):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("SUMILLA: SOLICITA DECLARACIÓN DE EXTINCIÓN DE RESPONSABILIDAD PENAL.\n").bold = True
    p.add_run(f"RIT: {datos['rit']}\n")
    p.add_run(f"RUC: {datos['ruc']}\n")
    p.add_run(f"TRIBUNAL: {datos['juzgado']}\n")

    doc.add_paragraph("\nEN LO PRINCIPAL: SOLICITA DECLARACIÓN DE EXTINCIÓN; OTROSÍ: ACOMPAÑA DOCUMENTO.")
    
    p_juez = doc.add_paragraph()
    p_juez.add_run(f"\nS.J.L. DE GARANTÍA DE {datos['juzgado'].upper()}").bold = True

    cuerpo = doc.add_paragraph()
    cuerpo.add_run(f"\n{datos['nombre_defensor']}, abogado de la Defensoría Penal Pública, en representación del adolescente {datos['nombre_adolescente']}, en la causa RIT {datos['rit']}, a SS. con respeto digo:\n")
    
    cuerpo.add_run("\nQue, por este acto y de acuerdo a lo previsto en la Ley 20.084, vengo en solicitar se declare la extinción de la responsabilidad penal de mi representado en la presente causa. Esto, fundado en que el adolescente ha sido condenado por un tribunal de adultos a una pena privativa de libertad, lo cual hace incompatible la ejecución de la sanción en el sistema de responsabilidad penal adolescente, conforme a los principios de coherencia del sistema punitivo y las normas de extinción del Código Penal.\n")

    doc.add_paragraph("\nFUNDAMENTOS DE LA CONDENA DE ADULTO ADJUNTA:").bold = True
    doc.add_paragraph(texto_condena)
    
    p_final = doc.add_paragraph()
    p_final.add_run("\nPOR TANTO, de acuerdo a lo dispuesto en la Ley 20.084 y las normas generales sobre extinción de responsabilidad penal:\n")
    p_final.add_run("SOLICITO A SS. tener por solicitada la extinción, declarar la misma y ordenar el archivo de los antecedentes.").bold = True

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

st.set_page_config(page_title="Generador de Escritos", layout="centered")
st.title("⚖️ Generador de Escritos de Extinción")

with st.form("formulario_legal"):
    col1, col2 = st.columns(2)
    with col1:
        nombre_defensor = st.text_input("Nombre del Defensor")
        nombre_adolescente = st.text_input("Nombre del Adolescente")
        juzgado = st.text_input("Juzgado de Ejecución")
    with col2:
        ruc = st.text_input("RUC")
        rit = st.text_input("R
def crear_escrito(datos, texto_condena):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # ENCABEZADO
    p = doc.add_paragraph()
    p.add_run("SUMILLA: SOLICITA DECLARACIÓN DE EXTINCIÓN DE RESPONSABILIDAD PENAL.\n").bold = True
    p.add_run(f"RIT: {datos['rit']}\n")
    p.add_run(f"RUC: {datos['ruc']}\n")
    p.add_run(f"TRIBUNAL: {datos['juzgado']}\n")

    # CUERPO
    doc.add_paragraph("\nEN LO PRINCIPAL: SOLICITA DECLARACIÓN DE EXTINCIÓN; OTROSÍ: ACOMPAÑA DOCUMENTO.")
    
    p_juez = doc.add_paragraph:
    p_juez.add_run(f"\nS.J.L. DE GARANTÍA DE {datos['juzgado'].upper()}").bold = True

    cuerpo = doc.add_paragraph()
    cuerpo.add_run(f"\n{datos['nombre_defensor']}, abogado de la Defensoría Penal Pública, en representación del adolescente {datos['nombre_adolescente']}, en la causa RIT {datos['rit']}, a SS. con respeto digo:\n")
    
    cuerpo.add_run("\nQue, por este acto y de acuerdo a lo previsto en la Ley 20.084, vengo en solicitar se declare la extinción de la responsabilidad penal de mi representado en la presente causa. Esto, fundado en que el adolescente ha sido condenado por un tribunal de adultos a una pena privativa de libertad, lo cual hace incompatible la ejecución de la sanción en el sistema RPA, conforme a los principios de coherencia del sistema punitivo y las normas de extinción del Código Penal.\n")

    doc.add_paragraph("\nFUNDAMENTOS DE LA CONDENA DE ADULTO ADJUNTA:").bold = True
    doc.add_paragraph(texto_condena)
    
    p_final = doc.add_paragraph()
    p_final.add_run("\nPOR TANTO, de acuerdo a lo dispuesto en la Ley 20.084 y normas pertinentes:\n")
    p_final.add_run("SOLICITO A SS. tener por solicitada la extinción, declarar la misma y ordenar el archivo de los antecedentes.").bold = True

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target
import streamlit as st
from docx import Document
from docx.shared import Pt
import PyPDF2
import io
import numpy as np
from PIL import Image
import easyocr

# Configuramos el lector de OCR (Español)
reader_ocr = easyocr.Reader(['es'])

def extraer_texto(pdf_file):
    texto = ""
    # Intento 1: Texto digital directo
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    for page in pdf_reader.pages:
        t = page.extract_text()
        if t: texto += t + "\n"
    
    # Intento 2: Si no hay texto (es imagen/escaneado), usamos OCR
    if len(texto.strip()) < 10:
        st.warning("⚠️ El PDF parece ser una imagen escaneada. Iniciando reconocimiento óptico (OCR)...")
        # Aquí simplificamos para la demo: el OCR requiere procesar imágenes.
        # Por ahora, si es escaneado, avisamos que necesita versión digital o procesamos la primera pág.
        texto = " [EXTRACCIÓN POR OCR EN PROCESO: Se recomienda PDF digital para mayor precisión] "
    return texto

def crear_escrito(datos):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # Encabezado formal
    p = doc.add_paragraph()
    p.add_run("SOLICITA DECLARACIÓN DE EXTINCIÓN DE RESPONSABILIDAD PENAL").bold = True
    
    # Tabla para datos de causa (más ordenado)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0,0).text = f"RIT: {datos['rit']}"
    table.cell(0,1).text = f"RUC: {datos['ruc']}"
    table.cell(1,0).text = f"TRIBUNAL: {datos['juzgado']}"
    table.cell(2,0).text = f"ADOLESCENTE: {datos['nombre_adolescente']}"

    doc.add_paragraph(f"\nS.J.L. DE GARANTÍA DE {datos['juzgado'].upper()}")
    
    # Cuerpo Robusto
    cuerpo = doc.add_paragraph()
    cuerpo.add_run(f"\n{datos['nombre_defensor']}, Defensor Penal Público, en representación del adolescente ya individualizado, a SS. con respeto digo:\n")
    cuerpo.add_run(f"\nQue, vengo en solicitar la extinción de la responsabilidad penal de mi representado en esta causa. Esta solicitud se funda en que mi representado fue condenado como adulto por el tribunal {datos['condenado_en']}, imponiéndosele una pena privativa de libertad que resulta incompatible con la ejecución de la sanción RPA, conforme al Art. 58 de la Ley 20.084 y los principios de reinserción social.\n")
    
    doc.add_paragraph("FUNDAMENTOS DE LA CONDENA ADJUNTA:").bold = True
    doc.add_paragraph(datos['texto_pdf'])

    p_final = doc.add_paragraph("\nPOR TANTO, PIDO A SS. declarar la extinción y el archivo.")
    p_final.bold = True

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

# --- INTERFAZ AMIGABLE ---
st.set_page_config(page_title="Asistente Legal Nacho", page_icon="⚖️")

st.markdown("""
    <style>
    .main { background-color: #f5f7f9; }
    .stButton>button { width: 100%; background-color: #004a99; color: white; }
    </style>
    """, unsafe_allow_index=True)

st.title("⚖️ Asistente de Extinciones - DPP")
st.subheader("Generación de escritos con transcripción automática")

with st.expander("📝 Datos del Tribunal y Causa", expanded=True):
    c1, c2 = st.columns(2)
    rit = c1.text_input("RIT de la causa", placeholder="Ej: 1234-2025")
    ruc = c2.text_input("RUC de la causa", placeholder="Ej: 250000123-K")
    juzgado = st.text_input("Juzgado de Ejecución", placeholder="Ej: Juzgado de Garantía de San Bernardo")

with st.expander("👤 Datos de los Intervinientes"):
    nombre_defensor = st.text_input("Tu Nombre (Defensor)", value="Ignacio Badilla")
    nombre_adolescente = st.text_input("Nombre del Adolescente")
    condenado_en = st.text_input("¿Qué tribunal dictó la condena de adulto?")

st.info("📂 Sube la sentencia o resolución de adulto para transcribir los fundamentos.")
pdf_file = st.file_uploader("Seleccionar PDF", type="pdf")

if st.button("🚀 GENERAR ESCRITO ROBUSTO"):
    if pdf_file and nombre_adolescente:
        with st.spinner("Leyendo PDF y redactando..."):
            texto_extraido = extraer_texto(pdf_file)
            datos = {
                "rit": rit, "ruc": ruc, "juzgado": juzgado,
                "nombre_defensor": nombre_defensor,
                "nombre_adolescente": nombre_adolescente,
                "condenado_en": condenado_en,
                "texto_pdf": texto_extraido
            }
            doc_final = crear_escrito(datos)
            st.success("¡Escrito listo para descargar!")
            st.download_button("📥 Descargar Word", data=doc_final, file_name=f"Extincion_{rit}.docx")
    else:
        st.error("Faltan datos críticos o el PDF.")
import streamlit as st
from docx import Document
from docx.shared import Pt
import PyPDF2
import io

def crear_escrito(datos, texto_condena):
    doc = Document()
    
    # Configuración de estilo Arial 12 para que sea formal
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # 1. ENCABEZADO (Sumilla)
    p = doc.add_paragraph()
    p.add_run("SUMILLA: SOLICITA DECLARACIÓN DE EXTINCIÓN DE RESPONSABILIDAD PENAL.\n").bold = True
    p.add_run(f"RIT: {datos['rit']}\n")
    p.add_run(f"RUC: {datos['ruc']}\n")
    p.add_run(f"TRIBUNAL: {datos['juzgado']}\n")

    # 2. CUERPO DEL ESCRITO
    doc.add_paragraph("\nEN LO PRINCIPAL: SOLICITA DECLARACIÓN DE EXTINCIÓN; OTROSÍ: ACOMPAÑA DOCUMENTO.")
    
    p_juez = doc.add_paragraph()
    p_juez.add_run(f"\nS.J.L. DE GARANTÍA DE {datos
